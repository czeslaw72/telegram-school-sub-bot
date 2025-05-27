import os
import re
import pandas as pd
import asyncio
import logging
from aiohttp import web
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackContext, CallbackQueryHandler
from telegram.ext.filters import BaseFilter, Document as TelegramDocument
from docx import Document

# Налаштування логування
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Отримуємо токен і порт із змінних середовища
TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "YOUR_TELEGRAM_BOT_TOKEN")
PORT = int(os.environ.get("PORT", 8000))

# Ініціалізація початкових даних
INITIAL_DATA = {
    "substitutions": {
        "Дата": ["02.05.2025", "02.05.2025"],
        "Клас": ["5", "6-А"],
        "Урок 0": ["інформатика", ""],
        "Урок 1": ["", "математика"],
        "Урок 2": ["", ""],
        "Урок 3": ["", ""],
        "Урок 4": ["", ""],
        "Урок 5": ["", ""],
        "Урок 6": ["", ""],
        "Урок 7": ["", ""]
    },
    "alert_mode": False,
    "users": []
}

# Зберігаємо стан у bot_data
substitutions_df = pd.DataFrame(INITIAL_DATA["substitutions"])
alert_mode = INITIAL_DATA["alert_mode"]
users = INITIAL_DATA["users"]
ADMIN_PASSWORD = "admin123"

# Перевірка пароля
def check_admin(password):
    return password == ADMIN_PASSWORD

# Очищення HTML-тегів із тексту
def clean_html_tags(text):
    if not isinstance(text, str):
        return ""
    cleaned = re.sub(r'<[^>]+>', '', text).strip()
    logger.info(f"Очищено текст: '{text}' -> '{cleaned}'")
    return cleaned

# Функція для витягнення таблиці з .docx
def extract_table_from_docx(file_path):
    try:
        logger.info("Починаємо зчитування .docx файлу")
        doc = Document(file_path)
        if not doc.tables:
            raise ValueError("У документі немає таблиць")
        
        table = doc.tables[0]
        data = []
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        logger.info(f"Отримані заголовки таблиці: {headers}")

        required_headers = ["Дата", "Клас"] + [f"Урок {i}" for i in range(0, 8)]
        if len(headers) != len(required_headers):
            raise ValueError(f"Очікується {len(required_headers)} стовпців, отримано {len(headers)}")
        missing_headers = [h for h in required_headers if h not in headers]
        if missing_headers:
            raise ValueError(f"Відсутні необхідні стовпці: {missing_headers}")

        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_data = [clean_html_tags(cell.text) for cell in row.cells]
            logger.info(f"Зчитано сирі дані рядка {row_idx}: {row_data}")
            if len(row_data) > len(headers):
                row_data = row_data[:len(headers)]
            while len(row_data) < len(headers):
                row_data.append("")
            logger.info(f"Нормалізовані дані рядка {row_idx}: {row_data}")
            if any(row_data):
                data.append(row_data)

        if not data:
            raise ValueError("Таблиця не містить даних після першого рядка")

        df = pd.DataFrame(data, columns=headers)
        logger.info(f"Зчитана таблиця:\n{df.to_string()}")
        return df
    except Exception as e:
        logger.error(f"Помилка при зчитуванні таблиці: {str(e)}")
        raise

# Функція для старту бота
async def start(update: Update, context: CallbackContext) -> None:
    global users
    user_id = update.effective_user.id
    if user_id not in users:
        users.append(user_id)
    context.bot_data.setdefault('substitutions_df', substitutions_df.copy())
    context.bot_data.setdefault('alert_mode', alert_mode)
    context.bot_data.setdefault('users', users.copy())

    keyboard = [
        [InlineKeyboardButton("Переглянути заміни", callback_data='view_subs')],
        [InlineKeyboardButton("Оновити таблицю (Адмін)", callback_data='update_subs')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Вітаю! Я бот для перегляду замін уроків. Оберіть опцію:", reply_markup=reply_markup)

# Обробка кнопок
async def button(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    await query.answer()

    if query.data == 'view_subs':
        keyboard = [[InlineKeyboardButton(cls, callback_data=f'class_{cls}')] for cls in context.bot_data['substitutions_df']["Клас"].unique()]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.message.reply_text("Оберіть клас:", reply_markup=reply_markup)
    elif query.data == 'update_subs':
        context.user_data['awaiting_password'] = True
        context.user_data['action'] = 'update_subs'
        await query.message.reply_text("Введіть пароль адміністратора:")

# Обробка вибору класу
async def handle_class_selection(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    await query.answer()
    class_name = query.data.split('_')[1]
    logger.info(f"Запит замін для класу: {class_name}")
    logger.info(f"Поточна таблиця замін перед відповіддю:\n{context.bot_data['substitutions_df'].to_string()}")
    class_subs = context.bot_data['substitutions_df'][context.bot_data['substitutions_df']["Клас"] == class_name]

    if class_subs.empty:
        await query.message.reply_text(f"Для класу {class_name} немає замін.")
        return

    response = f"Заміни для {class_name} на {class_subs['Дата'].iloc[0]}:\n"
    for col in [f"Урок {i}" for i in range(0, 8)]:
        value = class_subs[col].iloc[0]
        if value and value != "" and value != "-":
            response += f"{col}: {value}\n"
    await query.message.reply_text(response)

# Обробка текстових повідомлень
async def handle_text(update: Update, context: CallbackContext) -> None:
    message_text = update.message.text if update.message and update.message.text else ""
    logger.info(f"Отримане повідомлення: '{message_text}'")

    if context.user_data.get('awaiting_password'):
        if check_admin(message_text):
            context.user_data['is_admin'] = True
            context.user_data.pop('awaiting_password')
            action = context.user_data.get('action')
            if action == 'update_subs':
                await update.message.reply_text("Пароль правильний! Надішліть .docx файл із таблицею.")
        else:
            await update.message.reply_text("Неправильний пароль! Спробуйте ще раз:")
        return

    if "які заміни" in message_text.lower():
        logger.info(f"Запит через текст: {message_text}")
        logger.info(f"Поточна таблиця замін перед відповіддю:\n{context.bot_data['substitutions_df'].to_string()}")
        for class_name in context.bot_data['substitutions_df']["Клас"].unique():
            if class_name.lower() in message_text.lower():
                class_subs = context.bot_data['substitutions_df'][context.bot_data['substitutions_df']["Клас"] == class_name]
                if class_subs.empty:
                    await update.message.reply_text(f"Для класу {class_name} немає замін.")
                    return
                response = f"Заміни для {class_name} на {class_subs['Дата'].iloc[0]}:\n"
                for col in [f"Урок {i}" for i in range(0, 8)]:
                    value = class_subs[col].iloc[0]
                    if value and value != "" and value != "-":
                        response += f"{col}: {value}\n"
                await update.message.reply_text(response)
                return
        await update.message.reply_text("Будь ласка, вкажіть клас (наприклад, 'Які заміни в 5?').")
    else:
        await update.message.reply_text("Я розумію запити типу 'Які заміни в 5?'. Спробуйте ще раз!")

# Обробка .docx файлів
async def handle_docx(update: Update, context: CallbackContext) -> None:
    if not context.user_data.get('is_admin'):
        await update.message.reply_text("Спочатку увійдіть у режим адміністратора.")
        return

    if update.message.document and update.message.document.file_name.endswith('.docx'):
        file = await update.message.document.get_file()
        try:
            file_path = await file.download_to_drive()
            logger.info(f"Завантажено файл: {file_path}")
            new_df = extract_table_from_docx(file_path)
            context.bot_data['substitutions_df'] = new_df.copy()
            logger.info(f"Оновлена таблиця замін:\n{context.bot_data['substitutions_df'].to_string()}")
            context.user_data.pop('is_admin')
            await update.message.reply_text("Таблицю успішно оновлено з .docx файлу!")
        except Exception as e:
            logger.error(f"Помилка при обробці .docx файлу: {str(e)}")
            await update.message.reply_text(f"Помилка: {str(e)}")
        finally:
            if 'file_path' in locals():
                os.remove(file_path)
                logger.info(f"Видалено тимчасовий файл: {file_path}")

# Обробник Webhook-запитів
async def webhook(request):
    app = request.app['telegram_app']
    update = Update.de_json(await request.json(), app.bot)
    logger.info(f"Отримане оновлення: {update.to_dict()}")
    await app.process_update(update)
    return web.Response()

# Налаштування Webhook
async def setup_webhook(app: Application, webhook_url: str):
    logger.info(f"Налаштування Webhook на URL: {webhook_url}")
    await app.bot.set_webhook(url=webhook_url)

# Головна функція
async def main():
    logger.info(f"TELEGRAM_BOT_TOKEN: {TOKEN}")
    logger.info(f"RENDER_EXTERNAL_HOSTNAME: {os.environ.get('RENDER_EXTERNAL_HOSTNAME', 'Не встановлено')}")
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CallbackQueryHandler(button, pattern='^(view_subs|update_subs)$'))
    app.add_handler(CallbackQueryHandler(handle_class_selection, pattern='^class_'))
    app.add_handler(MessageHandler(BaseFilter(), handle_text))
    app.add_handler(MessageHandler(TelegramDocument(), handle_docx))

    web_app = web.Application()
    web_app['telegram_app'] = app
    web_app.router.add_post('/webhook', webhook)

    webhook_url = f"https://{os.environ.get('RENDER_EXTERNAL_HOSTNAME', 'your-service.onrender.com')}/webhook"
    await setup_webhook(app, webhook_url)

    await app.initialize()
    await app.start()

    runner = web.AppRunner(web_app)
    await runner.setup()
    site = web.TCPSite(runner, '0.0.0.0', PORT)
    await site.start()
    logger.info(f"Сервер запущено на порту {PORT}")

    await asyncio.Event().wait()

if __name__ == '__main__':
    asyncio.run(main())
